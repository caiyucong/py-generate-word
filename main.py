# encoding:UTF-8
import argparse
import json
from datetime import datetime

import redis
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches


def main(obj):
    result = init_data(obj)
    document = Document(args.template)
    init_doc(document)
    # 生成：1-7月份全市新签约项目总体情况表
    generate_paragraph(document, 16, WD_ALIGN_PARAGRAPH.CENTER, '1-7月份全市新签约项目总体情况表')
    generate_paragraph(document, 12, WD_ALIGN_PARAGRAPH.RIGHT, '单位：个')
    generate_table(document, '项目\n总体情况', ['', '新签约\n项目', '其中：\n制造业', '其中：\n服务业', '总投资\n（亿元）'],
                   result['sign_project_overview_list'], obj['signProjectOverviewList'],
                   ['projectCount', 'projectMadeCount', 'projectServiceCount', 'projectInvestAmountCount'])
    generate_table(document, '重大项目',
                   ['项目数', '其中：\n制造业', '其中：\n服务业', '其中：固投\n20-50亿元', '其中：固投\n50亿元以上'],
                   result['sign_project_major_list'], obj['signProjectMajorList'],
                   ['projectMadeCount', 'projectServiceCount', 'projectBasicInvestAmountTwentyCount',
                    'projectBasicInvestAmountFiftyCount'])
    # 增加段落 相当于换行
    document.add_paragraph()
    generate_table(document, '优强项目',
                   ['投资主体', '境内外\n500强\n含子公司', '上市公司', '独角兽或瞪羚企业', '专精特新企业\n（国家级）'],
                   result['sign_project_excellent_company_list'], obj['signProjectExcellentCompanyList'],
                   ['fiveHundredTopCompanyCount', 'listedCompanyCount', 'gazelleTagCompanyCount',
                    'specializedTagCompanyCount'])
    document.add_paragraph()
    generate_table(document, '项目来源地', ['', '个数', '个数占比', '总投资\n（亿元）', '投资占比'],
                   result['sign_project_county_list'], obj['signProjectCountyList'],
                   ['projectCount', 'projectCountPer%', 'projectAmountCount', 'projectAmountPer%'])
    generate_remark(document, content='1、新签约项目含制造业亿元以上投资项目、服务业2000万元以上投资项目；')
    generate_remark(document,
                    content='2、制造业重大项目为固投5亿元以上投资项目，服务业重大项目为总投资5000万元以上投资项目；')
    generate_remark(document,
                    content='3、境内外500强企业（世界500强、中国500强、民营500强）含子公司，上市公司、独角兽企业、瞪羚企业、专精特新企业（国家级）为直投。')
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    # 生成：新签约重大项目
    paragraph, _ = generate_paragraph(document, 16, content='二、新签约重大项目')
    first_indent(paragraph)
    paragraph, _ = generate_paragraph(document, 16,
                                      content='全市新签约重大项目42个（固投5亿元以上制造业项目31个，总投资5000万元以上服务业项目11个）。其中，固投20亿元以上项目8个，固投50'
                                              '亿元以上项目2 个（博望区宝明科技复合铜箔生产基地项目、市经开区正奇20GW高效N型电池片智能制造产业化项目）。')
    first_indent(paragraph)
    left_item = result['sign_major_project_county_list']
    document.add_paragraph()
    document.add_paragraph()
    generate_signed_major_table(document, left_item, obj['signMajorProjectCountyList'])
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    # 生成：亿元以上新开工项目
    paragraph, _ = generate_paragraph(document, 16, content='三、亿元以上新开工项目')
    first_indent(paragraph)
    paragraph, _ = generate_paragraph(document, 16, content='全市亿元以上新开工项目210个，完成全年目标任务93.3%。	'
                                                            '其中，20亿元以上项目17个，50亿元以上项目6'
                                                            '个（含山县爱柯迪新能源汽车零部件智能制造项目、和县天能特种电池生产项目、'
                                                            '当涂县新太高端合金新材料项目、雨山区国网国际绿色再制造项目、'
                                                            '市经开区中南高科马鞍山创智科技园项目、市经开区正奇20GW高效N型电池片智能制造产业化项目）。'
                                                            '新开工项目纳统106个，纳统率50.5%。')
    first_indent(paragraph)
    document.add_paragraph()
    document.add_paragraph()
    generate_table_2r_4c(document, ['新开工项目', '纳统率'], ['开工数（个）', '完成进度', '纳统数（个）', '纳统率'],
                         left_item,
                         '载体\\指标', ['projectCount', 'projectCountCompletePer%', 'synchronizationCount',
                                        'synchronizationCountPer%'], obj['startProjectCountyList'])
    document.add_paragraph()
    document.add_paragraph()
    # 生成：新投产固投2000万元以上制造业项目
    paragraph, _ = generate_paragraph(document, 16, content='四、新投产固投2000万元以上制造业项目')
    first_indent(paragraph)
    paragraph, _ = generate_paragraph(document, 16,
                                      content='全市新投产固投2000万元以上制造业项目166个，完成全年目标任务的91.7%。')
    first_indent(paragraph)
    document.add_paragraph()
    document.add_paragraph()
    generate_table_2000w(document, left_item, obj['operateProjectCountyList'])
    document.add_paragraph()
    document.add_paragraph()
    # 保存word文件
    document.save(args.filename)


def clone_paragraph_style(original_paragraph, new_paragraph):
    new_paragraph.style = original_paragraph.style
    for run in original_paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic


# 初始化文档
def init_doc(doc):
    current_date = datetime.now()
    formatted_date = current_date.strftime('%Y年%m月%d日')
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "${date}" in cell.text:
                    cell.text = cell.text.replace("${date}", formatted_date)
                    cell_alignment(cell)


# 首列缩进
def first_indent(paragraph, indent=0.5):
    paragraph.paragraph_format.first_line_indent = Inches(indent)


# 生成备注
def generate_remark(doc, font_size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT, content=''):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.alignment = alignment
    # 设置左间距
    paragraph.paragraph_format.left_indent = Inches(0.5)
    # 设置行间距
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(content)
    run.font.size = Pt(font_size)


# 生成段落
def generate_paragraph(doc, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, content=''):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.alignment = alignment
    run = paragraph.add_run(content)
    run.font.size = Pt(font_size)
    return paragraph, run


# 单元格居中对齐
def cell_alignment(cell):
    # 垂直居中
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 左右居中
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# 生成新签约重大项目
def generate_signed_major_table(doc, left_name=None, table_data=None, style='Table Grid'):
    if table_data is None:
        table_data = []
    if left_name is None:
        left_name = []
    row = 3 + len(left_name)
    column = 7
    table = doc.add_table(row, column, style=style)
    table.rows[0].height = Inches(0.4)
    table.rows[1].height = Inches(0.4)
    # 将第一列合并到第三列
    left_top = table.cell(0, 0)
    left_top.merge(table.cell(2, 0))
    left_top.text = '载体\\指标'
    cell_alignment(left_top)
    # 合并头标题
    header_title = table.cell(0, 1)
    header_title.merge(table.cell(0, column - 1))
    header_title.text = '新签约重大项目'
    cell_alignment(header_title)
    singe_number = table.cell(1, 1)
    singe_number.merge(table.cell(2, 1))
    singe_number.text = '签约数（个）'
    cell_alignment(singe_number)
    schedule = table.cell(1, 2)
    schedule.merge(table.cell(2, 2))
    schedule.text = '完成进度'
    cell_alignment(schedule)
    child_row_title = table.cell(1, 3)
    child_row_title.merge(table.cell(1, 5))
    child_row_title.text = '其中：制造业（个）'
    cell_alignment(child_row_title)
    child_5 = table.cell(2, 3)
    child_5.text = '固投\n5亿元\n以上'
    cell_alignment(child_5)
    child_20 = table.cell(2, 4)
    child_20.text = '固投\n20亿元\n以上'
    cell_alignment(child_20)
    child_50 = table.cell(2, 5)
    child_50.text = '固投\n50亿元\n以上'
    cell_alignment(child_50)
    cell_5000w = table.cell(1, 6)
    cell_5000w.merge(table.cell(2, 6))
    cell_5000w.text = '其中：5000万元以上服务业（个）'
    # 生成行标题
    for index, value in enumerate(left_name):
        table_row = table.rows[index + 3]
        table_row.height = Inches(0.4)
        item = table_row.cells[0]
        item.text = value
        item.width = Inches(1.5)
        cell_alignment(item)
    column_list = ['projectCount', 'projectCountCompletePer', 'fiveBasicAmountCount', 'twentyBasicAmountCount',
                   'fiftyBasicAmountCount', 'zeroFiveAmountCount']
    for index, obj in enumerate(table_data):
        item = table.rows[index + 3]
        for i, o in enumerate(column_list):
            e = item.cells[1 + i]
            e.text = str(obj[o])
            if o == 'projectCountCompletePer':
                e.text += '%'
            cell_alignment(e)


def generate_table_2000w(doc, left_title=None, table_data=None, style='Table Grid'):
    if table_data is None:
        table_data = []
    if left_title is None:
        left_title = []
    row = 2 + len(left_title)
    column = 3
    table = doc.add_table(row, column, style=style)
    left_top = table.cell(0, 0)
    left_top.text = '载体\\指标'
    cell_alignment(left_top)
    left_top.merge(table.cell(1, 0))
    main_title = table.cell(0, 1)
    main_title.merge(table.cell(0, 2))
    main_title.text = '新投产固投2000万元以上制造业项目'
    cell_alignment(main_title)
    number = table.cell(1, 1)
    number.text = '投产数（个）'
    cell_alignment(number)
    schedule = table.cell(1, 2)
    schedule.text = '完成进度'
    cell_alignment(schedule)
    for index, value in enumerate(left_title):
        item = table.rows[index + 2].cells[0]
        item.text = value
        cell_alignment(item)
    column_list = ['projectCount', 'projectCountCompletePer%']
    init_table_data(2, 1, table_data, column_list, table)
    set_row_height(table, 0.6)


# 生成第一行两列，第二行四列的表格
def generate_table_2r_4c(doc, header_title=None, child_header_title=None, left_title=None, left_top="",
                         column_name=None, table_data=None, style='Table Grid'):
    if table_data is None:
        table_data = []
    if column_name is None:
        column_name = []
    if left_title is None:
        left_title = []
    if child_header_title is None:
        child_header_title = []
    if header_title is None:
        header_title = []
    column = 5
    row = 2 + len(left_title)
    table = doc.add_table(row, column, style=style)
    left_top_cell = table.cell(0, 0)
    left_top_cell.text = left_top
    left_top_cell.merge(table.cell(1, 0))
    cell_alignment(left_top_cell)
    header_title_row = table.rows[0]
    header_title1 = header_title_row.cells[1]
    header_title1.merge(header_title_row.cells[2])
    header_title2 = header_title_row.cells[3]
    header_title2.merge(header_title_row.cells[4])
    header_title1.text = header_title[0]
    header_title2.text = header_title[1]
    cell_alignment(header_title1)
    cell_alignment(header_title2)
    set_row_height(table)
    child_header_title_rows = table.rows[1]
    for index, value in enumerate(child_header_title):
        item = child_header_title_rows.cells[index + 1]
        item.text = value
        cell_alignment(item)
    for index, value in enumerate(left_title):
        item = table.rows[index + 2].cells[0]
        item.text = value
        cell_alignment(item)
    init_table_data(2, 1, table_data, column_name, table)


def init_table_data(rows, column, table_data, column_name, table):
    for index, obj in enumerate(table_data):
        item = table.rows[index + rows]
        for i, o in enumerate(column_name):
            e = item.cells[column + i]
            if '%' in o:
                o = o.replace('%', '')
                cell_text = f'{str(obj[o])}%'
            else:
                cell_text = str(obj[o])
            e.text = cell_text
            cell_alignment(e)


# 生成表格
def generate_table(doc, left_text='', header=None, left_item=None, table_data=None, column_name=None,
                   style='Table Grid'):
    if column_name is None:
        column_name = []
    if table_data is None:
        table_data = []
    if left_item is None:
        left_item = []
    if header is None:
        header = []
    if len(column_name) != 0 and len(header) - 1 != len(column_name):
        return
    # 添加表格
    row = len(left_item) + 1
    column = len(header) + 1
    table = doc.add_table(row, column, style=style)
    # 获取左边列表
    left_content = table.cell(0, 0)
    left_content.text = left_text
    cell_alignment(left_content)
    # 合同单元格
    left_content.merge(table.cell(row - 1, 0))
    # 生成列头
    header_row = table.rows[0]
    for index, value in enumerate(header):
        item = header_row.cells[index + 1]
        item.text = value
        cell_alignment(item)
    # 生成行标题
    for index, value in enumerate(left_item):
        item = table.rows[index + 1].cells[1]
        item.text = str(value)
        cell_alignment(item)
    init_table_data(1, 2, table_data, column_name, table)
    set_row_height(table)
    return table


# 设置行高
def set_row_height(table, height=0.4):
    rows = table.rows
    for item in rows:
        item.height = Inches(height)


# 将字典的某个字段转为
def column_to_list(column, arr):
    new_arr = []
    for item in arr:
        new_arr.append(item[column])
    return new_arr


# 初始化数据 获取做最左列的值
def init_data(obj):
    sign_project_overview_list = column_to_list('projectCategory', obj['signProjectOverviewList'])
    sign_project_major_list = column_to_list('projectCount', obj['signProjectMajorList'])
    sign_project_excellent_company_list = column_to_list('excellentProjectCount',
                                                         obj['signProjectExcellentCompanyList'])
    sign_project_county_list = column_to_list('area', obj['signProjectCountyList'])
    sign_major_project_county_list = column_to_list('county', obj['signMajorProjectCountyList'])
    return {
        "sign_project_overview_list": sign_project_overview_list,
        "sign_project_major_list": sign_project_major_list,
        "sign_project_excellent_company_list": sign_project_excellent_company_list,
        "sign_major_project_county_list": sign_major_project_county_list,
        "sign_project_county_list": sign_project_county_list
    }


# 初始化reid
def init_redis_data(key):
    r = redis.StrictRedis(host=args.host, port=args.port, decode_responses=True, password=args.password)
    value = r.get(key)
    json_str = value.replace("\\", "")
    return json.loads(json_str[1:-1])


if __name__ == '__main__':
    # 创建 ArgumentParser 对象
    parser = argparse.ArgumentParser(description='Description of your program', add_help=False)
    # 添加参数
    parser.add_argument('-h', '--host', default='localhost', type=str, help='redis host')
    parser.add_argument('-p', '--port', default=6379, type=int, help='redis port')
    parser.add_argument('-pw', '--password', default=None, type=str, help='redis password')
    parser.add_argument('-k', '--key', default=None, type=str, help='redis key')
    parser.add_argument('-t', '--template', default=None, type=str, help='template 模板的路径')
    parser.add_argument('-n', '--filename', default=None, type=str, help='filename 保存的文件路径')
    # 解析命令行参数
    args = parser.parse_args()
    if args.key is not None:
        # 获取redis中的值
        val = init_redis_data(args.key)
        main(val)
        print('success')
